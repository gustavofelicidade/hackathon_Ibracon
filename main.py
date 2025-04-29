#!/usr/bin/env python3
"""
ESG.AUDITOR – Gera relatórios de sustentabilidade no padrão NBC TDS 01 / IFRS S1
Autor: Equipe UFRJ Analytica – Hackathon Ibracon 2025
"""

import os
from datetime import datetime

import openai           # pip install openai
import streamlit as st   # pip install streamlit
from docx import Document   # pip install python-docx

# ────────────────────────────────────────────────────────────────────────────────
# Configuração
# ────────────────────────────────────────────────────────────────────────────────
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
if not OPENAI_API_KEY:
    st.error(
        "⚠️ Defina a variável de ambiente OPENAI_API_KEY "
        "nas configurações do Railway antes de executar."
    )
    st.stop()

openai.api_key = OPENAI_API_KEY
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")  # altere se necessário

# ────────────────────────────────────────────────────────────────────────────────
# UI
# ────────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="ESG.AUDITOR", layout="wide")
st.title("📄 ESG.AUDITOR")
st.markdown(
    "Ferramenta simples e inteligente para gerar relatórios de sustentabilidade "
    "no formato exigido pelas normas **NBC TDS 01** e **IFRS S1**."
)

with st.form("esg_form", clear_on_submit=False):
    st.subheader("Informações da organização")
    company_name = st.text_input("Nome da empresa auditada")
    fiscal_year  = st.text_input("Exercício (AAAA)", value=str(datetime.now().year))

    st.subheader("Seções da norma")
    governanca = st.text_area("1 – Governança (estrutura, papéis e responsabilidades)")
    estrategia = st.text_area("2 – Estratégia ESG (objetivos, políticas e integração)")
    riscos     = st.text_area("3 – Riscos e Oportunidades (identificação e gestão)")
    metricas   = st.text_area("4 – Métricas e Metas (KPIs, indicadores quantitativos/qualitativos)")

    submitted = st.form_submit_button("🚀 Gerar relatório")

# ────────────────────────────────────────────────────────────────────────────────
# Geração do relatório
# ────────────────────────────────────────────────────────────────────────────────
if submitted:
    with st.spinner("Gerando relatório …"):
        user_content = f"""
Empresa: {company_name}
Exercício: {fiscal_year}

### Governança
{governanca}

### Estratégia ESG
{estrategia}

### Riscos e Oportunidades
{riscos}

### Métricas e Metas
{metricas}
        """

        system_content = (
            "Você é um auditor independente especialista em sustentabilidade. "
            "Com base nas respostas fornecidas, redija um relatório completo de sustentabilidade "
            "seguindo a estrutura NBC TDS 01 / IFRS S1, em linguagem formal e impessoal, "
            "mantendo os títulos de cada seção. "
            "Inclua comentários profissionais de asseguração quando pertinente."
        )

        # Chamada ao LLM
        response = openai.ChatCompletion.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_content},
                {"role": "user",   "content": user_content},
            ],
            temperature=0.3,
            max_tokens=2048,
        )
        report_md = response.choices[0].message.content

        # Converte para DOCX
        doc = Document()
        doc.add_heading(
            f"Relatório de Sustentabilidade – {company_name} – Exercício {fiscal_year}",
            level=1,
        )
        for line in report_md.split("\n"):
            if line.strip():
                if line.startswith("#"):
                    # Converte títulos markdown em heading 2
                    doc.add_heading(line.replace("#", "").strip(), level=2)
                else:
                    doc.add_paragraph(line.strip())

        filename = f"Relatorio_ESG_{company_name}_{fiscal_year}.docx".replace(" ", "_")
        doc.save(filename)

    st.success("✅ Relatório gerado com sucesso!")
    with open(filename, "rb") as f:
        st.download_button(
            label="💾 Baixar relatório (.docx)",
            data=f,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
